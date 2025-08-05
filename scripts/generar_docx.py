
import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def formatear_titulo(doc: Document, texto: str, nivel=2):
    p = doc.add_paragraph()
    run = p.add_run(texto.replace('_', ' ').capitalize())
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def formatear_contenido(doc: Document, texto: str):
    doc.add_paragraph(texto.strip())

def formatear_lista_guiones(doc: Document, items: list):
    for item in items:
        if isinstance(item, dict):
            texto = " - ".join(f"{v}" for v in item.values() if v)
        else:
            texto = str(item)
        doc.add_paragraph(f"- {texto}")

def formatear_cuantia(doc: Document, cuantias: list):
    if not cuantias:
        return
    max_len = max(len(item.get("concepto", "")) for item in cuantias if item.get("concepto"))
    for item in cuantias:
        concepto = item.get("concepto", "").strip()
        valor = item.get("valor", "").strip()
        unidad = item.get("unidad", "").strip()
        if concepto and valor:
            puntos = '.' * (max(5, 55 - len(concepto)))
            linea = f"- {concepto} {puntos} {valor} {unidad}".strip()
            doc.add_paragraph(linea)

def formatear_importe_maximo(doc: Document, items: list):
    for item in items:
        concepto = item.get("concepto", "").strip()
        cantidad = item.get("cantidad", "").strip()
        if concepto and cantidad:
            doc.add_paragraph(f"- {concepto} - {cantidad}")

def generar_docx_desde_json(json_path: str, output_dir: str = "salidas_docx"):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    nombre_base = os.path.splitext(os.path.basename(json_path))[0]
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{nombre_base}.docx")

    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except PermissionError:
            print(f"❌ No se puede sobrescribir {output_path}. Asegúrate de cerrar el archivo en Word.")
            return

    doc = Document()
    doc.add_heading("FICHA DE AYUDA UNIFICADA", level=1)

    for clave, valor in data.items():
        if valor in [None, "", [], "- "]:
            continue

        clave_limpia = clave.replace("_", " ").capitalize()

        if clave == "cuantia" and isinstance(valor, list):
            formatear_titulo(doc, clave_limpia)
            formatear_cuantia(doc, valor)

        elif clave == "importe_maximo" and isinstance(valor, list):
            formatear_titulo(doc, clave_limpia)
            formatear_importe_maximo(doc, valor)

        elif clave == "documentos_presentar" and isinstance(valor, list):
            formatear_titulo(doc, "Documentación a presentar")
            formatear_lista_guiones(doc, valor)

        elif isinstance(valor, str):
            formatear_titulo(doc, clave_limpia)
            formatear_contenido(doc, valor)

        elif isinstance(valor, list):
            if all(isinstance(v, str) for v in valor) or all(isinstance(v, dict) for v in valor):
                formatear_titulo(doc, clave_limpia)
                formatear_lista_guiones(doc, valor)

        elif isinstance(valor, dict):
            formatear_titulo(doc, clave_limpia)
            for subclave, subvalores in valor.items():
                if subvalores:
                    formatear_titulo(doc, subclave.capitalize())
                    formatear_lista_guiones(doc, subvalores)

    doc.save(output_path)
    print(f"📄 Documento Word generado en: {output_path}")
