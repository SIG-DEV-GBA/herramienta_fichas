import json
from docx import Document
from docx.shared import Pt
from pathlib import Path

def exportar_ficha_legible_a_docx(ruta_json: str, ruta_salida: str = "ficha_legible.docx"):
    """
    Genera un .docx legible desde una ficha JSON con formato legal.
    """
    ruta_json = Path(ruta_json)
    if not ruta_json.exists():
        print(f"❌ Archivo no encontrado: {ruta_json}")
        return

    with open(ruta_json, "r", encoding="utf-8") as f:
        ficha = json.load(f)

    doc = Document()
    doc.add_heading('Ficha Legal Generada', level=1)

    for campo, valor in ficha.items():
        doc.add_heading(campo.replace("_", " ").title(), level=2)

        # Si es lista o diccionario, lo mostramos como JSON legible
        if isinstance(valor, (dict, list)):
            texto = json.dumps(valor, indent=2, ensure_ascii=False)
        else:
            texto = str(valor)

        p = doc.add_paragraph()
        run = p.add_run(texto.strip())
        run.font.size = Pt(10)

    # Guardar documento
    ruta_docx = Path(ruta_salida)
    doc.save(ruta_docx)
    print(f"✅ Documento generado: {ruta_docx.resolve()}")

# Ejemplo de uso:
# exportar_ficha_legible_a_docx("salidas/ficha_documento_001.json")
